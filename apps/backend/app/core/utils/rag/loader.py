import io
import os
from typing import Callable

import boto3
import fitz
from docx import Document as DocxDocument
from langchain_core.documents import Document

from app.config import Config
from app.schemas.rag import DocumentData, Metadata
from app.database.models.documents import ContentType


EXTENSION_MAP: dict[str, str] = {
    ".pdf": "pdf",
    ".md": "md",
    ".docx": "docx",
    ".txt": "txt",
    ".pptx": "pptx",
    ".xlsx": "xlsx",
}

class DocumentLoader:
  # download content from s3
  # extract the type of file
  # load the document

  def __init__(self) -> None:
    self._s3 = self._connect_s3_client()
    self._bucket = Config["Env"].aws_bucket
    self._extractors: dict[
      ContentType, Callable[[bytes, Metadata], list[DocumentData]]
    ] = {
      ContentType.PDF: self._extract_pdf
    }

  def load(self, key: str) -> list[Document]:
    content_type = self._extract_type(key)
    file_bytes = self._download(key)

    metadata = Metadata(
      source=key,
      content_type=content_type,
      file_size=len(file_bytes),
      file_name=os.path.basename(key),
      page_count=file_bytes.count(b"\f") + 1 if content_type == ContentType.PDF else None
    )

    extractor = self._extractors.get(content_type)
    if not extractor:
      raise ValueError(f"No extractor for content type: {content_type}")

    extracted_docs = extractor(file_bytes, metadata)

    docs: list[Document] = []
    for item in extracted_docs:
      docs.append(Document(
        page_content=item.text,
        metadata = {**item.metadata.model_dump()}
      ))

    return docs

  def _connect_s3_client(self):
    return boto3.client(
        "s3", 
        aws_access_key_id = Config["Env"].aws_access_key, 
        aws_secret_access_key= Config["Env"].aws_secret_key,
        region_name=Config["Env"].aws_region
      )

  def _extract_type(self, key: str) -> ContentType:
    for ext, content_type in EXTENSION_MAP.items():
      if key.lower().endswith(ext):
        return ContentType(content_type)
    raise ValueError(f"Unsupported file type for key: {key}")

  def _download(self, key: str) -> bytes:
    response = self._s3.get_object(Bucket=self._bucket, Key=key)
    return response["Body"].read()

  def _extract_pdf(self, data: bytes, metadata: Metadata) -> list[DocumentData]:
    pages: list[DocumentData] = []
    doc = fitz.open(stream=data, filetype=metadata.content_type)
    metadata.page_count = len(doc)
    for page_index in range(metadata.page_count):
      page = doc.load_page(page_index)
      text = str(page.get_text("text"))
      page_metadata = Metadata(
        file_name=metadata.file_name,
        page_count=metadata.page_count,
        page=page_index + 1,
        source=metadata.source,
        content_type=metadata.content_type,
        char_count=len(text),
        file_size=metadata.file_size,
        word_count=len(text.split()) 
      )
      pages.append(DocumentData(text=text, metadata=page_metadata))
    doc.close()
    return pages
  
  def _extract_docx(self, data: bytes, metadata: Metadata) -> list[DocumentData]:
    doc = DocxDocument(io.BytesIO(data))
    texts: list[str] = []
    for d in doc.paragraphs:
      if d.text.strip():
        texts.append(d.text)
    text = "\n".join(texts)
    metadata.page_count = 1
    return [DocumentData(
      text=text,
      metadata=Metadata(
        char_count=len(text),
        content_type=metadata.content_type,
        file_name=metadata.file_name,
        file_size=metadata.file_size,
        page=metadata.page,
        page_count=metadata.page_count,
        source=metadata.source,
        word_count=len(text.split()),
      )
    )]
  
loader = DocumentLoader()
